import os
import yt_dlp
import ollama
from docx import Document

# ================= CẤU HÌNH CƠ BẢN =================
CHANNEL_URL = "https://www.youtube.com/@TheHiddenFiles" # Link channel cần cào
MAX_VIDEOS = 5 # Giới hạn số video chạy thử
OUTPUT_FILE = "KetQua_PhanTich_Raw.docx"
COOKIES_FILE = "cookies.txt"  
# ====================================================

def get_video_ids_from_channel(channel_url, max_videos=50):
    """Lấy danh sách Video ID từ Channel URL"""
    print(f"🔍 Đang quét danh sách video từ: {channel_url}")
    ydl_opts = {
        'extract_flat': True,
        'quiet': True,
        'playlistend': max_videos
    }
    
    video_ids = []
    try:
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(channel_url, download=False)
            if 'entries' in info:
                for entry in info['entries']:
                    if entry.get('id'):
                        video_ids.append(entry['id'])
    except Exception as e:
        print(f"❌ Lỗi khi lấy danh sách video: {e}")
        
    print(f"✅ Đã tìm thấy {len(video_ids)} video.")
    return video_ids

def fetch_caption(video_url):
    """Lấy caption dùng yt-dlp + Cookies"""
    ydl_opts = {
        'skip_download': True,
        'quiet': True,
        'no_warnings': True,
        'ignore_no_formats_error': True,
    }
    
    if os.path.exists(COOKIES_FILE):
        ydl_opts['cookiefile'] = COOKIES_FILE
    
    try:
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(video_url, download=False)
            
        if not info:
            return None
            
        subs = info.get('subtitles', {})
        auto_subs = info.get('automatic_captions', {})
        sub_url = None
        
        target_langs = [k for k in subs.keys() if k.lower().startswith('en')]
        if target_langs:
            preferred = 'en' if 'en' in target_langs else target_langs[0]
            for f in subs[preferred]:
                if f.get('ext') in ['json3', 'json']:
                    sub_url = f['url']
                    break
                    
        if not sub_url:
            for lang in auto_subs.keys():
                if lang.startswith('en'):  
                    for f in auto_subs[lang]:
                        if f.get('ext') in ['json3', 'json']:
                            sub_url = f['url']
                            break
                    if sub_url: break

        if not sub_url:
            return None
            
        try:
            import json
            with yt_dlp.YoutubeDL(ydl_opts) as ydl2:
                sub_data = ydl2.urlopen(sub_url).read().decode('utf-8')
            res = json.loads(sub_data)
        except Exception as e:
            if "429" in str(e): return "IP_BLOCKED"
            return None
            
        text_chunks = []
        for event in res.get('events', []):
            if 'segs' in event:
                for seg in event['segs']:
                    if 'utf8' in seg:
                        text_chunks.append(seg['utf8'])
                        
        full_text = "".join(text_chunks).replace('\n', ' ')
        # Giữ nguyên độ dài lớn để truyền 1 lần
        return ' '.join(full_text.split())[:150000] 
        
    except Exception as e:
        if "429" in str(e): return "IP_BLOCKED"
        return None

def analyze_caption_with_ai(caption_text):
    """Gửi TOÀN BỘ text thô cho Ollama trong 1 lần duy nhất"""
    if not caption_text or len(str(caption_text).strip()) < 50:
        return ""

    print(f"    -> Đang nạp toàn bộ văn bản ({len(caption_text)} ký tự) vào AI...")

    prompt = f"""Create a table with two columns: one in English and one in Vietnamese. Copy only the narrator's introduction from the article I provided, excluding the subtitles of the original video and timestamps. Correct spelling errors and capitalize the first letter of each sentence. Sentences should be separated into lines: continuous sentences (not interrupted by subtitles) should be on the same line, and sentences interrupted by subtitles should be on a separate line. Do not include the subtitles of the original video. Below is the text with all timestamps removed, retaining the original content:

{caption_text}"""

    try:
        response = ollama.chat(
            model='llama3.2:3b', 
            messages=[{'role': 'user', 'content': prompt}],
            options={
                'num_ctx': 32768, # Tăng giới hạn bộ nhớ ngữ cảnh lên rất cao để AI không bị quên chữ
                'temperature': 0.1
            }
        )
        
        return response['message']['content'].strip()
            
    except Exception as e:
        print(f"      ❌ Lỗi Llama AI: {e}")
        return ""

def create_word_document(all_data, output_filename):
    """Xuất văn bản thuần ra file Word"""
    print(f"\n💾 Đang xuất dữ liệu ra file {output_filename}...")
    doc = Document()
    
    doc.add_heading('Kết Quả Phân Tích Raw Từ AI', 0)

    for idx, text in enumerate(all_data, 1):
        if text.strip():
            doc.add_heading(f'Video {idx}', level=1)
            doc.add_paragraph(text)
            doc.add_paragraph("-" * 50) 

    doc.save(output_filename)
    print(f"🎉 Hoàn tất! Đã lưu kết quả thô vào file {output_filename}.")

def main():
    video_ids = get_video_ids_from_channel(CHANNEL_URL, MAX_VIDEOS)
    
    all_extracted_data = []
    
    for idx, vid in enumerate(video_ids, 1):
        video_url = f"https://www.youtube.com/watch?v={vid}"
        print(f"\n▶ Đang xử lý video {idx}/{len(video_ids)} (URL: {video_url})...")
        
        caption = fetch_caption(video_url)
        
        if caption == "IP_BLOCKED":
            print("🛑 DỪNG CHƯƠNG TRÌNH: IP đã bị YouTube chặn (Lỗi 429).")
            break
        if not caption:
            print("  -> Lỗi tải caption hoặc không có tiếng Anh. Bỏ qua.")
            continue
            
        # Ném toàn bộ caption vào AI 1 lần
        raw_response = analyze_caption_with_ai(caption)
        
        if raw_response:
            print(f"  -> ✅ Phân tích xong.")
            all_extracted_data.append(raw_response)
        else:
            print("  -> Lỗi AI hoặc trả về rỗng.")

    if all_extracted_data:
        create_word_document(all_extracted_data, OUTPUT_FILE)
    else:
        print("\nKhông có dữ liệu để lưu.")

if __name__ == "__main__":
    main()